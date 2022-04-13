# models.py file, for handeling and processing the input files

from asyncio.windows_events import NULL
from django.db import models
import hashlib
from PIL import Image
import pytesseract
import docx
import glob,os
import unicodedata
from fpdf import FPDF
import fitz
import random 
import string

# to generate random reference id for the input file
def random_value_generator(size=10, chars=string.ascii_uppercase+string.digits):
    return ''.join(random.choice(chars) for _ in range(size))

# interacting class with the database, to perform search function to return the retrived information
class FileManager(models.Manager):

    def search(self, query):
        return self.get_queryset().filter(models.Q(internal_reference__icontains=query) |
                                          models.Q(name__icontains=query) |
                                          models.Q(description__icontains=query)
                                          )

# main class function, that will be used to process the uploaded file
class File(models.Model):

    # information required for database upload
    name = models.CharField("Name", max_length=100)
    internal_reference = models.CharField("Internal Reference", max_length=100, editable=False)
    description = models.TextField("Description", blank=True, null=True)
    file = models.FileField(upload_to="OCR_file/input/", verbose_name="Input File")
    create_at = models.DateTimeField("Create at", auto_now_add=True)
    updated_at = models.DateTimeField("Update at", auto_now=True)

    # for correct (in required type) formatting
    def __str__(self):
        return "{0:03d} - {1}".format(self.id, self.file)

    # to extract text from image and save it into a .txt file
    def text_extraction(self,file):
        # opening the image
        img = Image.open(file)
        # extracting text
        txt = pytesseract.image_to_string(img, lang='eng')

        # saving content into .txt file
        my_file = open('static/text.txt', "w+")
        s=''
        for i in txt:
            # processing to remove null characters
            if(32<=ord(i)<=126 or ord(i)==10):
                s+=i

        # closing the file
        my_file.write(s)
        my_file.close()
        
        return my_file

    # creating doc file from .txt file
    def txt_to_doc(self):

        tfile=open('static/text.txt','r')
        # saving the content of txt file
        s=tfile.read()
        tfile.close()

        doc=docx.Document()
        # writing the content of txt file
        doc.add_paragraph(s)
        doc.save("static/text.docx")

        return doc

    # creating pdf file from .txt file
    def txt_to_pdf(self):
        
        tfile=open('static/text.txt','r')
        # saving the content of txt file
        s=tfile.read()
        tfile.close()
    
        # creating pdf
        pdf=FPDF('P','mm',(203,237))
        pdf.add_page()
        # setting default text style and size along with paper margins
        pdf.set_font('Arial','', 11)
        pdf.multi_cell(w=0,h=10,txt=s)
        # uploading the content of txt file
        pdf.output(name='static/text.pdf',dest='F')

        return pdf

    # for handeling use case of pdf upload
    def pdf_to_doc(self):
        mydoc = docx.Document() # document type

        # iterating the pdf's images, page wise
        with fitz.open('media/'+str(self.file)) as doc:
            txt=""
            for page in doc:
                txt+=page.get_text()
        
        # handeling the null characters of text read     
        s=''
        for i in txt:
            if(32<=ord(i)<=126 or ord(i)==10):
                s+=i

        # save as doc file    
        doc=docx.Document()
        doc.add_paragraph(txt)
        doc.save('static/text.docx')

        # save as pdf file
        pdf=FPDF('P','mm',(203,237))
        pdf.add_page()
        pdf.set_font('Arial','', 11)
        pdf.multi_cell(w=0,h=10,txt=s)
        pdf.output(name='static/text.pdf',dest='F')
        
        return mydoc


    # for handeling use case of image upload
    def execute_img_and_save_ocr(self):
        
        txtfile=self.text_extraction(self.file)
        docfile=self.txt_to_doc()
        pdffile=self.txt_to_pdf()
        # url links for doc file and pdf file returned
        return docfile,pdffile

    # storing the result to database
    def save(self, *args, **kwargs):

        if not self.internal_reference:
            # generating refernece id
            random_value = random_value_generator(size=20)

            while File.objects.filter(internal_reference=random_value).exists():
                random_value = random_value_generator(size=20)
            hash_value = hashlib.md5(bytes(str(self.id) + str(random_value), 'utf-8'))
            self.internal_reference = hash_value.hexdigest()
        # *args and **kwargs having the row properties to be added to table
        super(File, self).save(*args, **kwargs)

    # class for object upload to database
    class Meta:
        verbose_name = "PDFFile"
        verbose_name_plural = "PDFFiles"
        ordering = ['id']

    objects = FileManager()